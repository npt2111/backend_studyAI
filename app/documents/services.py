import base64
import gc
import io
import json as _json
import logging
import time
import re
import unicodedata
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional
from uuid import uuid4

from django.conf import settings
from docx import Document
from google import genai
from google.genai import types
from pypdf import PdfReader, PdfWriter
import requests

from config.services import supabase_client

logger = logging.getLogger(__name__)

_EASYOCR_READER = None


CHUNK_SYSTEM_PROMPT = """
Ban la tro ly tom tat hoc thuat tieng Viet. Nhan mot doan van ban, la mot phan cua tai lieu dai.

Nhiem vu: tom tat day du noi dung doan nay duoi dang JSON hop le.

Quy tac bat buoc:
1. Chi dung thong tin co trong doan van ban dau vao, khong suy dien, khong them kien thuc ngoai.
2. Giu nguyen ten chuong, muc, so thu tu neu co, vi du: "Chuong 3", "Muc 2.1".
3. Neu doan input thuc su rong hoac khong doc duoc, tra ve JSON voi `chapter_summary: "[THIEU_DU_LIEU]"`.
4. Khong viet them bat ky text nao ngoai JSON.
5. Khong boc JSON trong markdown code block.
6. Doc ky toan bo doan van, khong bo qua bang bieu, so lieu, dinh nghia quan trong.

Cau truc JSON dau ra:
{
  "chapters": [
    {
      "chapter_number": "1",
      "chapter_title": "Ten chuong hoac null neu khong co",
      "chapter_summary": "Tom tat toan bo noi dung chuong nay trong 2-4 cau, bam sat nguon. Bao gom so lieu, ket luan chinh neu co.",
      "sections": [
        {
          "section_number": "1.1",
          "section_title": "Ten muc",
          "section_summary": "Tom tat noi dung muc nay trong 1-3 cau, bam sat nguon."
        }
      ]
    }
  ],
  "key_points": [
    "Y chinh 1, 1-2 cau, bam sat nguon.",
    "Y chinh 2, 1-2 cau, bam sat nguon."
  ],
  "unclear_parts": "Ghi ro neu co doan bi cat, thieu, loi font; nguoc lai de chuoi rong."
}

Luu y quan trong:
- Neu doan input khong co cau truc chuong muc ro rang, tao 1 chapter voi `chapter_number: "0"`, `chapter_title: null`, `sections: []`.
- Mang `sections` co the rong neu khong co muc con.
- Trich 3-8 `key_points` tu doan nay, uu tien dinh nghia, so lieu, ket luan cot loi.
- Khong tra ve `[THIEU_DU_LIEU]` neu doan input co noi dung hop le, du ngan.
""".strip()


FINAL_SYSTEM_PROMPT = """
Ban la tro ly tong hop tom tat hoc thuat tieng Viet.
Ban nhan mot mang JSON, moi phan tu la ket qua tom tat cua mot doan trong cung mot tai lieu.

Nhiem vu: hop nhat tat ca chunks thanh mot ban tom tat hoan chinh, khong bo sot chuong muc nao.

Quy tac bat buoc:
1. Chi tong hop tu noi dung da cho, khong them kien thuc ngoai, khong suy dien.
2. Hop nhat cac chuong muc trung so thu tu tu cac chunks khac nhau cua cung mot chuong.
3. Giu nguyen so thu tu va ten chuong muc goc, sap xep theo thu tu tang dan.
4. Khong viet text nao ngoai JSON.
5. Khong boc JSON trong markdown code block.
6. Bao phu day du tat ca chuong muc xuat hien trong bat ky chunk nao.

Cau truc JSON dau ra:
{
  "chapters": [
    {
      "chapter_number": "1",
      "chapter_title": "Ten chuong hoac null",
      "chapter_summary": "Tom tat day du chuong nay trong 3-5 cau, bao quat toan bo noi dung, bao gom so lieu va ket luan chinh.",
      "sections": [
        {
          "section_number": "1.1",
          "section_title": "Ten muc",
          "section_summary": "Tom tat noi dung muc, 1-3 cau, bam sat nguon."
        }
      ]
    }
  ],
  "key_points": [
    "Y chinh quan trong nhat, 12-24 diem, moi diem 1-2 cau, bam sat nguon."
  ],
  "unclear_sections": [
    "Liet ke cac phan bi thieu, loi, cat ngan neu co; de mang rong neu khong co."
  ]
}

Huong dan key_points:
- 12-24 diem, uu tien y mang tinh ket luan, dinh nghia, so lieu, phuong phap cot loi.
- Moi diem la mot cau hoan chinh, co the dung doc lap, khong viet tat.

Huong dan hop nhat chapters:
- Neu nhieu chunks deu co "Chuong 2", gop tat ca sections va mo rong chapter_summary.
- Loai bo trung lap, giu thong tin day du nhat tu moi chunk.
""".strip()


KEYPOINTS_SYSTEM_PROMPT = """
Trich xuat key_points tu JSON tom tat dau vao.

Quy tac:
- Chi dung thong tin co trong dau vao, khong suy dien.
- Khong viet text nao ngoai JSON.
- Khong boc JSON trong markdown code block.
- Khong tra ve [THIEU_DU_LIEU] neu dau vao co noi dung hop le.

JSON dau ra:
{
  "key_points": [
    "Y chinh 1, cau hoan chinh, bam sat nguon.",
    "Y chinh 2, cau hoan chinh, bam sat nguon."
  ]
}

Tra ve 12-24 key_points.
""".strip()


# 
# Helpers
# 

SHORT_CHUNK_SYSTEM_PROMPT = """
Tom tat doan tai lieu thanh JSON thuan.
Chi dung thong tin trong input, khong suy dien, khong them kien thuc ngoai.
Giu ten chuong, muc, so thu tu neu co. Khong viet gi ngoai JSON, khong markdown.
Neu input rong hoac khong doc duoc, cho `chapter_summary = "[THIEU_DU_LIEU]"`.
Schema:
{"chapters":[{"chapter_number":"1","chapter_title":"Ten chuong hoac null","chapter_summary":"2-4 cau","sections":[{"section_number":"1.1","section_title":"Ten muc","section_summary":"1-3 cau"}]}],"key_points":["3-8 y chinh"],"unclear_parts":"chuoi rong hoac mo ta loi"}
Neu khong co cau truc chuong muc, tao 1 chapter voi `chapter_number="0"`, `chapter_title=null`, `sections=[]`.
""".strip()

SHORT_FINAL_SYSTEM_PROMPT = """
Hop nhat mang JSON cac chunk thanh 1 JSON tong hop day du, khong bo sot chuong muc.
Chi dung du lieu da cho, khong suy dien, khong viet gi ngoai JSON, khong markdown.
Gop cac chapter/section trung so thu tu, giu ten goc, sap xep tang dan.
Schema:
{"chapters":[{"chapter_number":"1","chapter_title":"Ten chuong hoac null","chapter_summary":"3-5 cau","sections":[{"section_number":"1.1","section_title":"Ten muc","section_summary":"1-3 cau"}]}],"key_points":["8-12 y chinh"],"unclear_sections":["cac phan mo ho neu co"]}
Key points la cau hoan chinh, bam sat nguon.
""".strip()

SHORT_KEYPOINTS_SYSTEM_PROMPT = """
Trich xuat key_points tu JSON dau vao.
Chi dung thong tin co san, khong suy dien, khong viet gi ngoai JSON, khong markdown.
Schema:
{"key_points":["8-12 cau y chinh"]}
""".strip()

CHUNK_SYSTEM_PROMPT = SHORT_CHUNK_SYSTEM_PROMPT
FINAL_SYSTEM_PROMPT = SHORT_FINAL_SYSTEM_PROMPT
KEYPOINTS_SYSTEM_PROMPT = SHORT_KEYPOINTS_SYSTEM_PROMPT

def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def normalize_job(row: Dict) -> Dict:
    if not row:
        return {}
    raw_points = row.get("key_points")
    key_points = raw_points if isinstance(raw_points, list) else []
    return {
        "id": row.get("id_job"),
        "id_job": row.get("id_job"),
        "user_id": row.get("id_user"),
        "file_name": row.get("file_name"),
        "status": row.get("status"),
        "progress": int(row.get("progress") or 0),
        "summary": row.get("summary_text"),
        "summary_text": row.get("summary_text"),
        "summary_json": row.get("summary_json"),
        "key_points": key_points,
        "error": row.get("error_message"),
        "created_at": row.get("created_at"),
        "updated_at": row.get("updated_at"),
    }


def _cleanup_text(raw_text: str) -> str:
    text = unicodedata.normalize("NFKC", raw_text or "")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("\u00ad", "").replace("\ufeff", "")
    text = re.sub(r"[^\S\n]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _docling_parse_text(
    file_name: str,
    mime_type: str,
    file_bytes: bytes,
    *,
    enable_ocr: Optional[bool] = None,
    force_full_page_ocr: bool = False,
) -> str:
    if not bool(getattr(settings, "DOCLING_ENABLED", True)):
        raise RuntimeError("Docling bi tat.")

    try:
        from docling.datamodel.base_models import InputFormat, DocumentStream
        from docling.datamodel.pipeline_options import (
            PdfPipelineOptions,
            TableStructureOptions,
            TesseractCliOcrOptions,
        )
        from docling.document_converter import DocumentConverter, PdfFormatOption
    except ImportError as exc:
        raise RuntimeError("Docling chua duoc cai dat.") from exc

    lower_name = (file_name or "").lower()
    lower_mime = (mime_type or "").lower()
    stream = DocumentStream(name=file_name or "document", stream=io.BytesIO(file_bytes))

    if lower_name.endswith(".pdf") or "pdf" in lower_mime:
        pipeline_options = PdfPipelineOptions()
        pipeline_options.do_table_structure = bool(getattr(settings, "DOCLING_TABLE_STRUCTURE", False))
        if pipeline_options.do_table_structure:
            pipeline_options.table_structure_options = TableStructureOptions(do_cell_matching=True)
        pipeline_options.do_ocr = (
            bool(getattr(settings, "DOCLING_DO_OCR", False))
            if enable_ocr is None
            else bool(enable_ocr)
        )
        if pipeline_options.do_ocr:
            ocr_options = TesseractCliOcrOptions(
                force_full_page_ocr=force_full_page_ocr
                or bool(getattr(settings, "DOCLING_FORCE_FULL_PAGE_OCR", False))
            )
            tesseract_cmd = str(getattr(settings, "TESSERACT_CMD", "") or "").strip()
            if tesseract_cmd:
                ocr_options.tesseract_cmd = tesseract_cmd
            pipeline_options.ocr_options = ocr_options
        converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options),
            }
        )
    else:
        converter = DocumentConverter()

    result = converter.convert(stream)
    document = result.document
    markdown = str(document.export_to_markdown() or "").strip()
    return _cleanup_text(markdown)


def _extract_pdf_text_with_mistral_ocr(file_bytes: bytes) -> str:
    if not bool(getattr(settings, "MISTRAL_OCR_ENABLED", False)):
        return ""

    api_key = str(getattr(settings, "MISTRAL_API_KEY", "") or "").strip()
    if not api_key:
        raise RuntimeError("MISTRAL_API_KEY chua duoc cau hinh.")

    encoded_pdf = base64.b64encode(file_bytes).decode("ascii")
    payload = {
        "model": str(getattr(settings, "MISTRAL_OCR_MODEL", "mistral-ocr-latest") or "mistral-ocr-latest"),
        "document": {
            "type": "document_url",
            "document_url": f"data:application/pdf;base64,{encoded_pdf}",
        },
        "include_image_base64": False,
    }
    timeout = int(getattr(settings, "MISTRAL_OCR_TIMEOUT_SECONDS", "180"))
    response = requests.post(
        "https://api.mistral.ai/v1/ocr",
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        json=payload,
        timeout=timeout,
    )
    if response.status_code >= 400:
        raise RuntimeError(f"Mistral OCR loi {response.status_code}: {response.text[:500]}")

    data = response.json()
    pages = data.get("pages") if isinstance(data, dict) else None
    if not isinstance(pages, list):
        raise RuntimeError("Mistral OCR khong tra ve danh sach pages hop le.")

    markdown_parts: List[str] = []
    for page in pages:
        if not isinstance(page, dict):
            continue
        markdown = str(page.get("markdown") or "").strip()
        if markdown:
            markdown_parts.append(markdown)

    text = _cleanup_text("\n\n".join(markdown_parts))
    if not text:
        raise RuntimeError("Mistral OCR khong trich xuat duoc noi dung PDF.")
    return text


def _extract_pdf_text(file_bytes: bytes) -> str:
    if bool(getattr(settings, "DOCLING_PDF_ENABLED", False)):
        try:
            text = _docling_parse_text(
                "document.pdf",
                "application/pdf",
                file_bytes,
                enable_ocr=False,
                force_full_page_ocr=False,
            )
            if text:
                return text
        except Exception as exc:
            logger.warning("Docling PDF parse failed; falling back to pypdf extract: %s", exc)

    reader = PdfReader(io.BytesIO(file_bytes))
    chunks = [(page.extract_text() or "").strip() for page in reader.pages]
    return "\n\n".join([c for c in chunks if c])


def _extract_pdf_text_layer_markdown(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    blocks: List[str] = []
    for idx, page in enumerate(reader.pages, start=1):
        page_text = _cleanup_text(page.extract_text() or "")
        if page_text:
            blocks.append(f"## Page {idx}\n\n{page_text}")
    return _cleanup_text("\n\n".join(blocks))


def _is_page_text_good(text: str) -> bool:
    min_words = int(getattr(settings, "PDFPLUMBER_TEXT_MIN_WORDS_PER_PAGE", "20"))
    words = text.split()
    if len(words) < min_words:
        return False
    alpha_count = sum(1 for ch in text if ch.isalpha())
    printable_count = sum(1 for ch in text if ch.isprintable())
    return printable_count > 0 and (alpha_count / max(printable_count, 1)) >= 0.20


def _easyocr_reader():
    global _EASYOCR_READER
    if _EASYOCR_READER is not None:
        return _EASYOCR_READER

    try:
        import easyocr
    except ImportError as exc:
        raise RuntimeError("EasyOCR chua duoc cai dat.") from exc

    langs = [
        item.strip()
        for item in str(getattr(settings, "EASYOCR_LANGS", "vi,en") or "vi,en").split(",")
        if item.strip()
    ]
    if not langs:
        langs = ["vi", "en"]
    _EASYOCR_READER = easyocr.Reader(langs, gpu=bool(getattr(settings, "EASYOCR_GPU", False)))
    return _EASYOCR_READER


def _resize_image_for_ocr(image):
    max_side = int(getattr(settings, "EASYOCR_MAX_IMAGE_SIDE", "1800"))
    if max_side <= 0:
        return image
    width, height = image.size
    longest = max(width, height)
    if longest <= max_side:
        return image
    scale = max_side / float(longest)
    new_size = (max(1, int(width * scale)), max(1, int(height * scale)))
    return image.resize(new_size)


def _ocr_pdfplumber_page(page) -> str:
    if not bool(getattr(settings, "EASYOCR_ENABLED", True)):
        return ""

    try:
        import numpy as np
    except ImportError as exc:
        raise RuntimeError("numpy chua duoc cai dat cho EasyOCR.") from exc

    dpi = int(getattr(settings, "EASYOCR_DPI", "130"))
    page_image = page.to_image(resolution=dpi)
    image = page_image.original.convert("RGB")
    image = _resize_image_for_ocr(image)
    image_array = np.array(image)

    reader = _easyocr_reader()
    results = reader.readtext(
        image_array,
        detail=0,
        paragraph=bool(getattr(settings, "EASYOCR_PARAGRAPH", False)),
        batch_size=1,
    )
    del image_array
    del image
    del page_image
    gc.collect()

    lines = [str(item).strip() for item in results if str(item).strip()]
    return _cleanup_text("\n".join(lines))


def _extract_pdf_markdown_with_pdfplumber_easyocr(file_bytes: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("pdfplumber chua duoc cai dat.") from exc

    blocks: List[str] = []
    ocr_pages: List[int] = []
    text_pages = 0
    max_pages = int(getattr(settings, "PDF_OCR_MAX_PAGES", "0"))

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        total_pages = len(pdf.pages)
        for idx, page in enumerate(pdf.pages, start=1):
            if max_pages > 0 and idx > max_pages:
                logger.warning("PDF OCR stopped at configured page limit %s/%s.", max_pages, total_pages)
                break

            page_text = _cleanup_text(page.extract_text() or "")
            source = "pdfplumber"
            if not _is_page_text_good(page_text):
                try:
                    page_text = _ocr_pdfplumber_page(page)
                    source = "easyocr"
                    if page_text:
                        ocr_pages.append(idx)
                except Exception as exc:
                    logger.warning("EasyOCR failed for PDF page %s: %s", idx, exc)
                    page_text = ""

            if page_text:
                text_pages += 1
                blocks.append(f"## Page {idx} ({source})\n\n{page_text}")

            if hasattr(page, "flush_cache"):
                page.flush_cache()
            gc.collect()

    text = _cleanup_text("\n\n".join(blocks))
    if text:
        logger.info("PDF parsed with pdfplumber/easyocr: text_pages=%s, ocr_pages=%s", text_pages, ocr_pages)
    return text


def _extract_pdf_markdown_with_pdfplumber_pymupdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("pdfplumber chua duoc cai dat.") from exc

    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError("PyMuPDF chua duoc cai dat.") from exc

    blocks: List[str] = []
    text_pages = 0
    pymupdf_pages: List[int] = []
    max_pages = int(getattr(settings, "PDF_OCR_MAX_PAGES", "0"))

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            total_pages = len(pdf.pages)
            for idx, page in enumerate(pdf.pages, start=1):
                if max_pages > 0 and idx > max_pages:
                    logger.warning("PDF text extraction stopped at configured page limit %s/%s.", max_pages, total_pages)
                    break

                page_text = _cleanup_text(page.extract_text() or "")
                source = "pdfplumber"
                if not _is_page_text_good(page_text) and idx - 1 < len(doc):
                    page_text = _cleanup_text(doc[idx - 1].get_text("text") or "")
                    source = "pymupdf"
                    if page_text:
                        pymupdf_pages.append(idx)

                if page_text:
                    text_pages += 1
                    blocks.append(f"## Page {idx} ({source})\n\n{page_text}")

                if hasattr(page, "flush_cache"):
                    page.flush_cache()
                gc.collect()
    finally:
        doc.close()

    text = _cleanup_text("\n\n".join(blocks))
    if text:
        logger.info("PDF parsed with pdfplumber/pymupdf: text_pages=%s, pymupdf_pages=%s", text_pages, pymupdf_pages)
    return text


def _extract_pdf_pages(file_bytes: bytes) -> List[str]:
    reader = PdfReader(io.BytesIO(file_bytes))
    return [(page.extract_text() or "").strip() for page in reader.pages]


def _is_pdf_text_extractable(pages: List[str], min_alpha_ratio: float = 0.25) -> bool:
    total_text = " ".join([p for p in pages if p])
    if len(total_text.split()) < 80:
        return False
    alpha = sum(1 for ch in total_text if ch.isalpha())
    printable = sum(1 for ch in total_text if ch.isprintable())
    return printable > 0 and (alpha / max(printable, 1)) >= min_alpha_ratio


def _extract_docx_text(file_bytes: bytes) -> str:
    try:
        text = _docling_parse_text(
            "document.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_bytes,
        )
        if text:
            return text
    except Exception as exc:
        logger.warning("Docling DOCX parse failed; falling back to python-docx: %s", exc)

    doc = Document(io.BytesIO(file_bytes))
    blocks: List[str] = []
    blocks.extend([p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()])
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks)


def _extract_text(file_name: str, mime_type: str, file_bytes: bytes) -> str:
    lower_name = (file_name or "").lower()
    lower_mime = (mime_type or "").lower()
    if lower_name.endswith(".pdf") or "pdf" in lower_mime:
        return _extract_pdf_text(file_bytes)
    if lower_name.endswith(".docx") or "wordprocessingml" in lower_mime:
        return _extract_docx_text(file_bytes)
    raise RuntimeError("Chi ho tro PDF va DOCX.")


def _extract_pdf_text_with_ocr(file_bytes: bytes) -> str:
    if not bool(getattr(settings, "DOCLING_DO_OCR", False)):
        return ""

    try:
        text = _docling_parse_text(
            "document.pdf",
            "application/pdf",
            file_bytes,
            enable_ocr=True,
            force_full_page_ocr=True,
        )
        if text:
            return text
    except Exception as exc:
        logger.warning("Docling PDF OCR parse failed; falling back to standard PDF parse: %s", exc)
    return ""


def _validate_document_file(file_name: str, mime_type: str, file_bytes: bytes) -> None:
    lower_name = (file_name or "").lower()
    lower_mime = (mime_type or "").lower()

    if lower_name.endswith(".pdf") or "pdf" in lower_mime:
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            if len(reader.pages) <= 0:
                raise RuntimeError("PDF khong co trang hop le.")
        except Exception as exc:
            raise RuntimeError("PDF khong hop le hoac bi hong.") from exc
        return

    if lower_name.endswith(".docx") or "wordprocessingml" in lower_mime:
        text = _cleanup_text(_extract_docx_text(file_bytes))
        if not text:
            raise RuntimeError("Khong trich xuat duoc noi dung file DOCX.")
        return

    raise RuntimeError("Chi ho tro PDF va DOCX.")


def _chunk_text(text: str, max_chars: int) -> List[str]:
    if len(text) <= max_chars:
        return [text]
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: List[str] = []
    current = ""
    for para in paragraphs:
        candidate = f"{current}\n\n{para}" if current else para
        if len(candidate) <= max_chars:
            current = candidate
        else:
            if current:
                chunks.append(current)
            if len(para) <= max_chars:
                current = para
            else:
                start = 0
                while start < len(para):
                    end = min(start + max_chars, len(para))
                    chunks.append(para[start:end])
                    start = end
                current = ""
    if current:
        chunks.append(current)
    return chunks


def _extract_outline_headings(text: str) -> List[str]:
    lines = [ln.strip() for ln in text.splitlines() if ln and ln.strip()]
    headings: List[str] = []
    patterns = [
        r"^(chuong|chng)\s+\d+[\.: -].*",
        r"^(muc|mc)\s+\d+[\.: -].*",
        r"^\d+(\.\d+){0,3}\s+.+",
        r"^[ivxlcdm]+\.\s+.+",
    ]
    for ln in lines:
        low = ln.lower()
        if any(re.match(p, low) for p in patterns):
            headings.append(ln)
    seen = set()
    result: List[str] = []
    for h in headings:
        key = h.lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(h)
    return result[:200]


def _chunk_text_by_headings(text: str, max_chars: int) -> List[str]:
    lines = text.splitlines()
    heading_re = re.compile(
        r"^((chuong|chng)\s+\d+[\.: -].*|(muc|mc)\s+\d+[\.: -].*|\d+(\.\d+){0,3}\s+.+|[ivxlcdm]+\.\s+.+)$",
        flags=re.IGNORECASE,
    )
    sections: List[List[str]] = []
    current: List[str] = []
    for ln in lines:
        stripped = ln.strip()
        if stripped and heading_re.match(stripped):
            if current:
                sections.append(current)
            current = [ln]
        else:
            if not current:
                current = [ln]
            else:
                current.append(ln)
    if current:
        sections.append(current)

    blocks = ["\n".join(sec).strip() for sec in sections if "\n".join(sec).strip()]
    if not blocks:
        return _chunk_text(text, max_chars=max_chars)

    chunks: List[str] = []
    current_chunk = ""
    for block in blocks:
        candidate = f"{current_chunk}\n\n{block}" if current_chunk else block
        if len(candidate) <= max_chars:
            current_chunk = candidate
        else:
            if current_chunk:
                chunks.append(current_chunk)
            if len(block) <= max_chars:
                current_chunk = block
            else:
                chunks.extend(_chunk_text(block, max_chars=max_chars))
                current_chunk = ""
    if current_chunk:
        chunks.append(current_chunk)
    return chunks if chunks else _chunk_text(text, max_chars=max_chars)


def _validate_source_text(text: str) -> None:
    words = text.split()
    if len(words) < 80:
        raise RuntimeError("Noi dung trich xuat qua ngan de tom tat day du.")
    alpha_count = sum(1 for ch in text if ch.isalpha())
    printable_count = sum(1 for ch in text if ch.isprintable())
    if printable_count == 0 or (alpha_count / max(printable_count, 1)) < 0.25:
        raise RuntimeError("Noi dung trich xuat chat luong thap (co the la PDF scan/loi font).")
    if text.count("\ufffd") > 10:
        raise RuntimeError("Noi dung trich xuat bi loi ky tu, khong the tom tat chinh xac.")


def _validate_pdf_markdown_text(text: str) -> None:
    min_words = int(getattr(settings, "MISTRAL_OCR_MIN_WORDS", "80"))
    words = text.split()
    if len(words) < min_words:
        raise RuntimeError("Khong doc duoc noi dung file.")

    alpha_count = sum(1 for ch in text if ch.isalpha())
    printable_count = sum(1 for ch in text if ch.isprintable())
    if printable_count == 0 or (alpha_count / max(printable_count, 1)) < 0.25:
        raise RuntimeError("Khong doc duoc noi dung file.")

    if text.count("\ufffd") > 10:
        raise RuntimeError("Khong doc duoc noi dung file.")


# 
# JSON parse / validate / repair helpers
# 

def _safe_parse_json(raw: str) -> Dict:
    """Strip markdown fences if the model returns ```json, then parse."""
    text = (raw or "").strip()
    # Remove ```json ... ``` if present.
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    text = text.strip()
    # Find the first JSON object in the response.
    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        raise RuntimeError("Khong tim thay JSON hop le trong response.")
    candidate = match.group(0).strip()
    try:
        return _json.loads(candidate)
    except _json.JSONDecodeError:
        repaired = re.sub(r",\s*([}\]])", r"\1", candidate)
        repaired = re.sub(r'(["\]\}0-9A-Za-z])\s*\n\s*(")', r"\1,\n\2", repaired)
        repaired = re.sub(r'(")\s+("[-A-Za-z0-9_]+\"\s*:)', r'\1, \2', repaired)
        return _json.loads(repaired)


def _validate_summary_json(data: Dict) -> None:
    """Validate final summary JSON structure."""
    if not isinstance(data, dict):
        raise RuntimeError("Response khong phi dict JSON.")
    chapters = data.get("chapters")
    if not isinstance(chapters, list) or len(chapters) == 0:
        raise RuntimeError("JSON thieu truong 'chapters' hoac rong.")
    key_points = data.get("key_points")
    if not isinstance(key_points, list) or len(key_points) < 3:
        raise RuntimeError("JSON thieu 'key_points' hoac qua it diem.")
    # Check excessive THIEU_DU_LIEU markers.
    raw_str = _json.dumps(data, ensure_ascii=False).lower()
    thieu_count = raw_str.count("thieu_du_lieu")
    if thieu_count > len(chapters) // 2:
        raise RuntimeError("Tom tat chua day du, qua nhieu [THIEU_DU_LIEU].")


def _merge_chunk_jsons(chunk_raws: List[str]) -> List[Dict]:
    """
    Parse each raw chunk string into a dict.
    If parsing fails, keep raw_text so the final step can still handle it.
    """
    result = []
    for raw in chunk_raws:
        try:
            result.append(_safe_parse_json(raw))
        except Exception:
            result.append({"raw_text": (raw or "")[:3000]})
    return result


def _fallback_summary_json_from_chunks(chunk_dicts: List[Dict]) -> Dict:
    chapters: List[Dict] = []
    key_points: List[str] = []
    unclear_sections: List[str] = []

    chapter_map: Dict[str, Dict] = {}
    section_seen: Dict[str, set] = {}

    for item in chunk_dicts:
        for ch in item.get("chapters", []) if isinstance(item, dict) else []:
            chapter_number = str(ch.get("chapter_number") or "0").strip()
            chapter_title = ch.get("chapter_title")
            chapter_key = chapter_number or "0"
            if chapter_key not in chapter_map:
                chapter_map[chapter_key] = {
                    "chapter_number": chapter_number or "0",
                    "chapter_title": chapter_title,
                    "chapter_summary": str(ch.get("chapter_summary") or "").strip(),
                    "sections": [],
                }
                section_seen[chapter_key] = set()
            else:
                current_summary = chapter_map[chapter_key].get("chapter_summary", "")
                new_summary = str(ch.get("chapter_summary") or "").strip()
                if len(new_summary) > len(current_summary):
                    chapter_map[chapter_key]["chapter_summary"] = new_summary
                if not chapter_map[chapter_key].get("chapter_title") and chapter_title:
                    chapter_map[chapter_key]["chapter_title"] = chapter_title

            for sec in ch.get("sections", []) if isinstance(ch, dict) else []:
                sec_num = str(sec.get("section_number") or "").strip()
                sec_title = str(sec.get("section_title") or "").strip()
                sec_key = f"{sec_num}|{sec_title}".lower()
                if sec_key in section_seen[chapter_key]:
                    continue
                section_seen[chapter_key].add(sec_key)
                chapter_map[chapter_key]["sections"].append(
                    {
                        "section_number": sec_num,
                        "section_title": sec.get("section_title"),
                        "section_summary": str(sec.get("section_summary") or "").strip(),
                    }
                )

        for kp in item.get("key_points", []) if isinstance(item, dict) else []:
            text = str(kp).strip()
            if text and text not in key_points:
                key_points.append(text)

        for unclear in item.get("unclear_sections", []) if isinstance(item, dict) else []:
            text = str(unclear).strip()
            if text and text not in unclear_sections:
                unclear_sections.append(text)

        raw_text = str(item.get("raw_text") or "").strip() if isinstance(item, dict) else ""
        if raw_text and raw_text not in unclear_sections:
            unclear_sections.append("Mot chunk khong parse duoc JSON tong hop.")

    chapters = sorted(
        chapter_map.values(),
        key=lambda x: (str(x.get("chapter_number") or "9999")),
    )

    if not chapters:
        chapters = [
            {
                "chapter_number": "0",
                "chapter_title": None,
                "chapter_summary": "Khong tong hop duoc day du tu ket qua chunk.",
                "sections": [],
            }
        ]

    return {
        "chapters": chapters,
        "key_points": key_points[:12],
        "unclear_sections": unclear_sections[:20],
    }


def _repair_summary_json(client: genai.Client, raw: str) -> Dict:
    """Repair malformed summary JSON by calling the model."""
    if not bool(getattr(settings, "SUMMARY_ENABLE_MODEL_REPAIR", False)):
        raise RuntimeError("Model repair is disabled.")
    fixed_raw = _chat(
        client=client,
        system_prompt=(
            "Convert the following content into valid JSON with the required schema. "
            "Do not change content. Do not add new knowledge. "
            "Return plain JSON only, no markdown and no extra text. "
            'Required schema: {"chapters": [...], "key_points": [...], "unclear_sections": []}'
        ),
        user_prompt=raw[:4000],
        max_tokens=int(getattr(settings, "SUMMARY_REPAIR_MAX_TOKENS", "900")),
    )
    return _safe_parse_json(fixed_raw)


def _parse_key_points_from_json(data: Dict) -> List[str]:
    """Get key_points from parsed summary JSON."""
    return [str(p).strip() for p in data.get("key_points", []) if str(p).strip()]


def _summary_json_to_text(data: Dict) -> str:
    """
    Convert summary JSON to plain text for legacy DB and simple UI rendering.
    """
    lines = []
    for ch in data.get("chapters", []):
        num = ch.get("chapter_number", "")
        title = ch.get("chapter_title") or ""
        header = f"Chuong {num}" if num and str(num) != "0" else ""
        if title:
            header = f"{header}: {title}" if header else title
        if header:
            lines.append(f"## {header}")
        if ch.get("chapter_summary"):
            lines.append(ch["chapter_summary"])
        for sec in ch.get("sections", []):
            sec_num = sec.get("section_number", "")
            sec_title = sec.get("section_title") or ""
            sec_header = f"Muc {sec_num}" if sec_num else ""
            if sec_title:
                sec_header = f"{sec_header}: {sec_title}" if sec_header else sec_title
            if sec_header:
                lines.append(f"### {sec_header}")
            if sec.get("section_summary"):
                lines.append(sec["section_summary"])
    if data.get("key_points"):
        lines.append("\n## Cac y chinh")
        lines.extend([f"- {p}" for p in data["key_points"]])
    return "\n\n".join(lines).strip()


# 
# Sanitize helpers
# 

def _sanitize_summary_text(summary_text: str) -> str:
    text = (summary_text or "").strip()
    text = text.replace("**", "")
    text = re.sub(r"\[THIEU_DU_LIEU\]", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\berror\b", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _sanitize_key_points(points: List[str]) -> List[str]:
    cleaned_points: List[str] = []
    seen = set()
    for p in points:
        t = _sanitize_summary_text(p)
        t = re.sub(
            r"^(TOM_TAT_THEO_MUC|CAC_Y_CHINH_KHONG_BO_SOT|NOI_DUNG_CHUA_RO)\s*:?\s*$",
            "", t, flags=re.IGNORECASE,
        )
        t = t.strip("- ").strip()
        if not t:
            continue
        if "thieu_du_lieu" in t.lower() or t.lower() == "error":
            continue
        key = t.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned_points.append(t)
    return cleaned_points[:12]


def _derive_key_points_from_summary_json(data: Dict) -> List[str]:
    points: List[str] = []
    seen = set()

    def add_point(value: str) -> None:
        text = _sanitize_summary_text(str(value or ""))
        text = re.sub(r"\s+", " ", text).strip("- ").strip()
        if not text:
            return
        if "thieu_du_lieu" in text.lower() or text.lower() == "error":
            return
        key = text.lower()
        if key in seen:
            return
        seen.add(key)
        points.append(text)

    for ch in data.get("chapters", []) if isinstance(data, dict) else []:
        if not isinstance(ch, dict):
            continue
        chapter_title = str(ch.get("chapter_title") or "").strip()
        chapter_summary = str(ch.get("chapter_summary") or "").strip()
        if chapter_summary:
            add_point(f"{chapter_title}: {chapter_summary}" if chapter_title else chapter_summary)

        for sec in ch.get("sections", []) if isinstance(ch, dict) else []:
            if not isinstance(sec, dict):
                continue
            section_title = str(sec.get("section_title") or "").strip()
            section_summary = str(sec.get("section_summary") or "").strip()
            if section_summary:
                add_point(f"{section_title}: {section_summary}" if section_title else section_summary)

    return points[:12]


def _derive_key_points_from_summary_text(summary_text: str) -> List[str]:
    text = _sanitize_summary_text(summary_text)
    if not text:
        return []

    candidates: List[str] = []
    for line in text.splitlines():
        cleaned = re.sub(r"^\s*[-*#\d\.\)\(]+\s*", "", line).strip()
        if len(cleaned.split()) >= 6:
            candidates.append(cleaned)

    if not candidates:
        candidates = [
            part.strip()
            for part in re.split(r"(?<=[.!?])\s+", text)
            if len(part.strip().split()) >= 6
        ]

    return _sanitize_key_points(candidates[:12])


def _sanitize_summary_json(data: Dict) -> Dict:
    """Sanitize summary JSON fields."""
    cleaned = dict(data)
    # Lm sch chapters
    chapters = []
    for ch in cleaned.get("chapters", []):
        ch_clean = dict(ch)
        ch_clean["chapter_summary"] = _sanitize_summary_text(ch.get("chapter_summary", ""))
        sections = []
        for sec in ch.get("sections", []):
            sec_clean = dict(sec)
            sec_clean["section_summary"] = _sanitize_summary_text(sec.get("section_summary", ""))
            if sec_clean.get("section_summary"):
                sections.append(sec_clean)
        ch_clean["sections"] = sections
        if ch_clean.get("chapter_summary") or sections:
            chapters.append(ch_clean)
    cleaned["chapters"] = chapters
    cleaned["key_points"] = _sanitize_key_points(cleaned.get("key_points", []))
    cleaned.pop("keywords", None)
    cleaned["unclear_sections"] = [str(s).strip() for s in cleaned.get("unclear_sections", []) if str(s).strip()]
    return cleaned


# 
# Coverage audit
# 

def _normalize_heading_for_match(s: str) -> str:
    s = unicodedata.normalize("NFKD", s.lower())
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"[^a-z0-9\s]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _coverage_audit(source_text: str, summary_text: str) -> Dict:
    source_headings = _extract_outline_headings(source_text)
    if not source_headings:
        return {
            "source_headings": [],
            "matched_headings": [],
            "missing_headings": [],
            "coverage_ratio": 1.0 if _sanitize_summary_text(summary_text) else 0.0,
        }
    summary_norm = _normalize_heading_for_match(summary_text)
    matched: List[str] = []
    missing: List[str] = []
    for h in source_headings:
        h_norm = _normalize_heading_for_match(h)
        if h_norm and h_norm in summary_norm:
            matched.append(h)
        else:
            missing.append(h)
    ratio = len(matched) / max(1, len(source_headings))
    return {
        "source_headings": source_headings,
        "matched_headings": matched,
        "missing_headings": missing,
        "coverage_ratio": ratio,
    }


# 
# Storage helpers
# 

def _build_summary_json_payload(
    *,
    job_id: str,
    file_name: str,
    summary_text: str,
    summary_json: Dict,
    key_points: List[str],
    source_word_count: int,
    coverage: Dict,
) -> Dict:
    return {
        "job_id": job_id,
        "file_name": file_name,
        "summary_text": summary_text,
        "summary_json": summary_json,
        "key_points": key_points,
        "source_word_count": source_word_count,
        "coverage": coverage,
        "generated_at": now_iso(),
    }


def _upload_summary_json(
    *,
    bucket: str,
    user_id: str,
    job_id: str,
    payload: Dict,
) -> str:
    object_path = f"{user_id}/summaries/{job_id}_{uuid4().hex}.json"
    file_bytes = _json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    res, status_code = supabase_client.upload_storage_file(
        bucket=bucket,
        object_path=object_path,
        file_bytes=file_bytes,
        content_type="application/json; charset=utf-8",
    )
    if status_code >= 400:
        raise RuntimeError(f"Khong luu duoc file JSON summary len Supabase Storage: {res}")
    return object_path


# 
# Gemini client & chat helper
# 

def _gemini_client() -> genai.Client:
    api_key = getattr(settings, "GEMINI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("GEMINI_API_KEY chua duoc cau hinh.")
    return genai.Client(api_key=api_key)


def _extract_gemini_text(response) -> str:
    content = str(getattr(response, "text", "") or "").strip()
    if content:
        return content
    texts: List[str] = []
    for candidate in getattr(response, "candidates", []) or []:
        parts = getattr(getattr(candidate, "content", None), "parts", None) or []
        for part in parts:
            part_text = getattr(part, "text", None)
            if part_text:
                texts.append(str(part_text).strip())
    return "\n".join([t for t in texts if t]).strip()


def _gemini_models_to_try() -> List[str]:
    primary = str(getattr(settings, "GEMINI_MODEL", "gemini-3.1-flash-lite") or "").strip()
    fallback = str(getattr(settings, "GEMINI_FALLBACK_MODEL", "gemini-2.5-flash") or "").strip()
    models: List[str] = []
    for model in [primary, fallback]:
        if model and model not in models:
            models.append(model)
    return models


def _should_try_next_gemini_model(exc: Exception) -> bool:
    msg = str(exc).lower()
    return (
        "not found" in msg
        or "not supported" in msg
        or "permission" in msg
        or "invalid argument" in msg
        or "400" in msg
        or "404" in msg
        or "503" in msg
        or "unavailable" in msg
        or "high demand" in msg
    )


def _is_retryable_gemini_error(exc: Exception) -> bool:
    msg = str(exc).lower()
    return (
        "rate_limit" in msg
        or "429" in msg
        or "too many" in msg
        or "resource_exhausted" in msg
        or "503" in msg
        or "unavailable" in msg
        or "high demand" in msg
    )


def _chat_groq(system_prompt: str, user_prompt: str, max_tokens: int) -> str:
    api_key = str(getattr(settings, "GROQ_API_KEY", "") or "").strip()
    if not api_key:
        raise RuntimeError("GROQ_API_KEY chua duoc cau hinh.")

    model = str(getattr(settings, "GROQ_MODEL", "llama-3.3-70b-versatile") or "").strip()
    max_retries = int(getattr(settings, "GROQ_RETRY_MAX", "3"))
    base_sleep = float(getattr(settings, "GROQ_RETRY_BASE_SECONDS", "8"))
    last_exc = None

    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.2,
        "max_tokens": max_tokens,
    }

    for attempt in range(max_retries):
        try:
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json=payload,
                timeout=120,
            )
            if response.status_code >= 400:
                raise RuntimeError(f"Groq loi {response.status_code}: {response.text[:500]}")
            data = response.json()
            choices = data.get("choices") if isinstance(data, dict) else None
            if not choices:
                raise RuntimeError("Groq tra ve noi dung rong.")
            message = choices[0].get("message") if isinstance(choices[0], dict) else None
            content = str((message or {}).get("content") or "").strip()
            if not content:
                raise RuntimeError("Groq tra ve noi dung rong.")
            return content
        except Exception as exc:
            last_exc = exc
            msg = str(exc).lower()
            if (
                "429" in msg
                or "rate" in msg
                or "too many" in msg
                or "503" in msg
                or "unavailable" in msg
                or "timeout" in msg
            ):
                if attempt < max_retries - 1:
                    time.sleep(base_sleep * (2 ** attempt))
                    continue
            raise

    if last_exc:
        raise last_exc
    raise RuntimeError("Groq: khong co response sau retry.")


def _chat_ollama(system_prompt: str, user_prompt: str, max_tokens: int) -> str:
    base_url = str(getattr(settings, "OLLAMA_BASE_URL", "http://localhost:11434") or "http://localhost:11434").rstrip("/")
    model = str(getattr(settings, "OLLAMA_MODEL", "llama3.2:1b") or "llama3.2:1b").strip()
    timeout = int(getattr(settings, "OLLAMA_TIMEOUT_SECONDS", "180"))

    payload = {
        "model": model,
        "stream": False,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "options": {
            "temperature": 0.2,
            "num_predict": max_tokens,
        },
    }
    response = requests.post(
        f"{base_url}/api/chat",
        json=payload,
        timeout=timeout,
    )
    if response.status_code >= 400:
        raise RuntimeError(f"Ollama loi {response.status_code}: {response.text[:500]}")

    data = response.json()
    message = data.get("message") if isinstance(data, dict) else None
    content = str((message or {}).get("content") or "").strip()
    if not content:
        content = str(data.get("response") or "").strip() if isinstance(data, dict) else ""
    if not content:
        raise RuntimeError("Ollama tra ve noi dung rong.")
    return content


def _chat(client: genai.Client, system_prompt: str, user_prompt: str, max_tokens: int) -> str:
    provider = str(getattr(settings, "SUMMARY_LLM_PROVIDER", "groq") or "groq").lower()
    if provider == "ollama":
        return _chat_ollama(system_prompt=system_prompt, user_prompt=user_prompt, max_tokens=max_tokens)
    if provider == "groq":
        return _chat_groq(system_prompt=system_prompt, user_prompt=user_prompt, max_tokens=max_tokens)

    max_retries = int(getattr(settings, "GEMINI_RETRY_MAX", "3"))
    base_sleep = float(getattr(settings, "GEMINI_RETRY_BASE_SECONDS", "8"))

    last_exc = None
    models = _gemini_models_to_try()
    for model_idx, model in enumerate(models):
        for attempt in range(max_retries):
            try:
                response = client.models.generate_content(
                    model=model,
                    contents=user_prompt,
                    config=types.GenerateContentConfig(
                        system_instruction=system_prompt,
                        temperature=0.2,
                        max_output_tokens=max_tokens,
                    ),
                )
                content = _extract_gemini_text(response)
                if not content:
                    raise RuntimeError("Gemini tra ve noi dung rong.")
                return content
            except Exception as exc:
                last_exc = exc
                msg = str(exc).lower()
                if _is_retryable_gemini_error(exc):
                    if attempt < max_retries - 1:
                        time.sleep(base_sleep * (2 ** attempt))
                        continue
                if model_idx < len(models) - 1 and _should_try_next_gemini_model(exc):
                    logger.warning("Gemini model %s failed, trying fallback model: %s", model, exc)
                    break
                raise

    if last_exc:
        raise last_exc
    raise RuntimeError("Gemini: khong co response sau retry.")


def _chat_with_document(
    client: genai.Client,
    system_prompt: str,
    user_prompt: str,
    file_bytes: bytes,
    mime_type: str,
    max_tokens: int,
) -> str:
    max_retries = int(getattr(settings, "GEMINI_RETRY_MAX", "3"))
    base_sleep = float(getattr(settings, "GEMINI_RETRY_BASE_SECONDS", "8"))

    last_exc = None
    models = _gemini_models_to_try()
    for model_idx, model in enumerate(models):
        for attempt in range(max_retries):
            try:
                response = client.models.generate_content(
                    model=model,
                    contents=[
                        types.Part.from_bytes(data=file_bytes, mime_type=mime_type),
                        user_prompt,
                    ],
                    config=types.GenerateContentConfig(
                        system_instruction=system_prompt,
                        temperature=0.2,
                        max_output_tokens=max_tokens,
                    ),
                )
                content = _extract_gemini_text(response)
                if not content:
                    raise RuntimeError("Gemini tra ve noi dung rong khi doc tai lieu.")
                return content
            except Exception as exc:
                last_exc = exc
                msg = str(exc).lower()
                if _is_retryable_gemini_error(exc):
                    if attempt < max_retries - 1:
                        time.sleep(base_sleep * (2 ** attempt))
                        continue
                if model_idx < len(models) - 1 and _should_try_next_gemini_model(exc):
                    logger.warning("Gemini document model %s failed, trying fallback model: %s", model, exc)
                    break
                raise

    if last_exc:
        raise last_exc
    raise RuntimeError("Gemini: khong co response sau retry khi doc tai lieu.")


# 
# PDF scan fallback  chia trang, extract text, tm tt tng batch
# 

def _summarize_pdf_pages_via_text(
    *,
    client: genai.Client,
    file_bytes: bytes,
    job_id: str,
    max_pages_per_chunk: int = 8,
) -> Dict:
    """
    Doc PDF bang Gemini native document input theo tung batch trang,
    sau do tong hop ve JSON cuoi cung.
    """
    reader = PdfReader(io.BytesIO(file_bytes))
    total_pages = len(reader.pages)
    if total_pages <= 0:
        raise RuntimeError("PDF khong co trang hop le.")

    page_groups: List[tuple] = []
    for start in range(0, total_pages, max_pages_per_chunk):
        end = min(start + max_pages_per_chunk, total_pages)
        page_groups.append((start, end))

    chunk_raws: List[str] = []
    total_groups = len(page_groups)

    for idx, (start, end) in enumerate(page_groups, start=1):
        writer = PdfWriter()
        for page_idx in range(start, end):
            writer.add_page(reader.pages[page_idx])

        buf = io.BytesIO()
        writer.write(buf)
        chunk_bytes = buf.getvalue()

        raw = _chat_with_document(
            client=client,
            system_prompt=CHUNK_SYSTEM_PROMPT,
            user_prompt=(
                f"[TRANG {start + 1}-{end}/{total_pages}] "
                "Doc day du cac trang PDF nay va tom tat thanh JSON theo schema. "
                "Khong bo sot bang bieu, so lieu, hinh ve, dinh nghia, heading."
            ),
            file_bytes=chunk_bytes,
            mime_type="application/pdf",
            max_tokens=int(getattr(settings, "SUMMARY_CHUNK_MAX_TOKENS", "650")),
        )
        chunk_raws.append(raw)

        progress = min(80, 20 + int((idx / total_groups) * 55))
        if idx == total_groups or idx % 2 == 0:
            supabase_client.update_summary_job(job_id, {"progress": progress})

    merged_chunks = _merge_chunk_jsons(chunk_raws)
    merged_input = _json.dumps(merged_chunks, ensure_ascii=False)

    final_raw = _chat(
        client=client,
        system_prompt=FINAL_SYSTEM_PROMPT,
        user_prompt=merged_input,
        max_tokens=int(getattr(settings, "SUMMARY_FINAL_MAX_TOKENS", "1200")),
    )
    try:
        summary_data = _safe_parse_json(final_raw)
    except Exception:
        summary_data = _fallback_summary_json_from_chunks(merged_chunks)

    try:
        _validate_summary_json(summary_data)
    except RuntimeError:
        if bool(getattr(settings, "SUMMARY_ENABLE_MODEL_REPAIR", False)):
            try:
                summary_data = _repair_summary_json(client, final_raw)
                _validate_summary_json(summary_data)
            except Exception:
                summary_data = _fallback_summary_json_from_chunks(merged_chunks)
        else:
            summary_data = _fallback_summary_json_from_chunks(merged_chunks)

    summary_data = _sanitize_summary_json(summary_data)
    summary_text = _summary_json_to_text(summary_data)

    all_pages_text = _extract_pdf_pages(file_bytes)
    source_text = _cleanup_text("\n\n".join(all_pages_text))
    coverage = (
        _coverage_audit(source_text, summary_text)
        if source_text
        else {
            "coverage_ratio": 0.0,
            "source_headings": [],
            "matched_headings": [],
            "missing_headings": [],
        }
    )
    return {
        "summary_data": summary_data,
        "summary_text": summary_text,
        "coverage": coverage,
        "source_text": source_text,
    }

# 
# Main summarize pipeline: chunk by headings/paragraphs, retry if coverage is low
# 

def _summarize_with_chunks_retry(
    *,
    client: genai.Client,
    text: str,
    job_id: str,
) -> Dict:
    max_chunk_chars    = int(getattr(settings, "SUMMARY_CHUNK_CHARS", 6000))
    coverage_threshold = float(getattr(settings, "SUMMARY_COVERAGE_THRESHOLD", "0.6"))

    attempt_plans = [
        {"max_chars": max_chunk_chars, "extra_prompt": ""},
        {
            "max_chars": max(3000, int(max_chunk_chars * 0.75)),
            "extra_prompt": "Tap trung bao phu day du tung chuong/muc, khong bo sot.",
        },
    ]
    max_attempts  = int(getattr(settings, "SUMMARY_RETRY_ATTEMPTS", "1"))
    attempt_plans = attempt_plans[: max(1, min(max_attempts, len(attempt_plans)))]

    best: Dict = {
        "summary_data": {},
        "summary_text": "",
        "coverage": {
            "coverage_ratio": 0.0,
            "source_headings": [],
            "matched_headings": [],
            "missing_headings": [],
        },
    }

    for attempt_idx, plan in enumerate(attempt_plans, start=1):
        chunks = _chunk_text_by_headings(text, max_chars=int(plan["max_chars"]))
        if not chunks:
            chunks = _chunk_text(text, max_chars=int(plan["max_chars"]))
        if not chunks:
            raise RuntimeError("Khong tach duoc chunk.")

        chunk_raws: List[str] = []
        total = len(chunks)

        for idx, chunk in enumerate(chunks, start=1):
            raw = _chat(
                client=client,
                system_prompt=CHUNK_SYSTEM_PROMPT,
                user_prompt=(
                    f"[PHAN {idx}/{total}] {plan['extra_prompt']}\n\n{chunk}"
                ),
                max_tokens=int(getattr(settings, "SUMMARY_CHUNK_MAX_TOKENS", "650")),
            )
            chunk_raws.append(raw)
            progress = min(85, 20 + int((idx / total) * 55))
            if idx == total or idx % 2 == 0:
                supabase_client.update_summary_job(job_id, {"progress": progress})

        # Merge chunk JSONs, then call final summary.
        merged_chunks = _merge_chunk_jsons(chunk_raws)
        merged_input  = _json.dumps(merged_chunks, ensure_ascii=False)

        final_raw = _chat(
            client=client,
            system_prompt=FINAL_SYSTEM_PROMPT,
            user_prompt=merged_input,
            max_tokens=int(getattr(settings, "SUMMARY_FINAL_MAX_TOKENS", "1200")),
        )
        try:
            summary_data = _safe_parse_json(final_raw)
        except Exception:
            summary_data = _fallback_summary_json_from_chunks(merged_chunks)

        try:
            _validate_summary_json(summary_data)
        except RuntimeError:
            if bool(getattr(settings, "SUMMARY_ENABLE_MODEL_REPAIR", False)):
                try:
                    summary_data = _repair_summary_json(client, final_raw)
                    _validate_summary_json(summary_data)
                except Exception:
                    summary_data = _fallback_summary_json_from_chunks(merged_chunks)
            else:
                summary_data = _fallback_summary_json_from_chunks(merged_chunks)

        summary_data = _sanitize_summary_json(summary_data)
        summary_text = _summary_json_to_text(summary_data)
        coverage     = _coverage_audit(text, summary_text)

        if coverage.get("coverage_ratio", 0.0) >= coverage_threshold:
            return {
                "summary_data": summary_data,
                "summary_text": summary_text,
                "coverage": coverage,
            }

        if summary_text and (
            not best.get("summary_text")
            or coverage.get("coverage_ratio", 0.0) >= best["coverage"]["coverage_ratio"]
        ):
            best = {
                "summary_data": summary_data,
                "summary_text": summary_text,
                "coverage": coverage,
            }

        if attempt_idx < len(attempt_plans):
            supabase_client.update_summary_job(job_id, {"progress": 60})

    return best


# 
# Main entry point
# 

def process_summary_job(job_id: str) -> None:
    claimed_row, claimed_status = supabase_client.claim_summary_job(job_id)
    if claimed_status >= 400:
        raise RuntimeError("Khong claim duoc job de xu ly.")
    if not claimed_row:
        return

    try:
        bucket = getattr(settings, "SUPABASE_STORAGE_BUCKET", "study-documents")

        # 1. Download file.
        blob, blob_status = supabase_client.download_storage_file(
            bucket=bucket,
            object_path=str(claimed_row.get("storage_path", "")),
        )
        if blob_status >= 400 or not isinstance(blob, (bytes, bytearray)):
            raise RuntimeError("Khong tai duoc file tu Supabase Storage.")

        file_name  = str(claimed_row.get("file_name", ""))
        user_id    = str(claimed_row.get("id_user", ""))
        mime_type  = str(claimed_row.get("mime_type", ""))
        lower_name = file_name.lower()
        lower_mime = mime_type.lower()
        is_pdf     = lower_name.endswith(".pdf") or ("pdf" in lower_mime)

        text = ""

        if not is_pdf:
            text = _extract_text(file_name, mime_type, bytes(blob))
            text = _cleanup_text(text)
            if not text:
                raise RuntimeError("Khong trich xuat duoc noi dung file.")

        supabase_client.update_summary_job(job_id, {"progress": 20})

        llm_provider = str(getattr(settings, "SUMMARY_LLM_PROVIDER", "groq") or "groq").lower()
        client: Optional[genai.Client] = None if llm_provider in {"groq", "ollama"} else _gemini_client()
        coverage: Dict = {
            "coverage_ratio": 0.0,
            "source_headings": [],
            "matched_headings": [],
            "missing_headings": [],
        }
        summary_data: Dict = {}
        final_summary_text: str = ""

        if is_pdf:
            supabase_client.update_summary_job(job_id, {"progress": 15})
            reader = str(getattr(settings, "SUMMARY_DOCUMENT_READER", "docling") or "docling").lower()
            if reader == "mistral":
                if not bool(getattr(settings, "MISTRAL_OCR_ENABLED", False)):
                    raise RuntimeError("Mistral OCR chua duoc bat cho PDF.")
                text = _cleanup_text(_extract_pdf_text_with_mistral_ocr(bytes(blob)))
                logger.info("Mistral OCR extracted PDF markdown for job %s.", job_id)
            elif reader in {"pdfplumber_pymupdf", "pymupdf"}:
                text = _extract_pdf_markdown_with_pdfplumber_pymupdf(bytes(blob))
                logger.info("pdfplumber/pymupdf extracted PDF markdown for job %s.", job_id)
            elif reader in {"pdfplumber", "pdfplumber_easyocr", "easyocr"}:
                text = _extract_pdf_markdown_with_pdfplumber_easyocr(bytes(blob))
                logger.info("pdfplumber/easyocr extracted PDF markdown for job %s.", job_id)
            else:
                if not bool(getattr(settings, "DOCLING_ENABLED", True)):
                    raise RuntimeError("Docling chua duoc bat.")
                try:
                    text = _cleanup_text(
                        _docling_parse_text(
                            file_name=file_name,
                            mime_type=mime_type,
                            file_bytes=bytes(blob),
                            enable_ocr=bool(getattr(settings, "DOCLING_DO_OCR", False)),
                            force_full_page_ocr=bool(getattr(settings, "DOCLING_FORCE_FULL_PAGE_OCR", False)),
                        )
                    )
                    logger.info("Docling extracted PDF markdown for job %s.", job_id)
                except Exception as exc:
                    if not bool(getattr(settings, "DOCLING_PDF_FALLBACK_TEXT_LAYER", True)):
                        raise
                    logger.warning("Docling PDF parse failed; falling back to PDF text layer: %s", exc)
                    text = _extract_pdf_text_layer_markdown(bytes(blob))
            _validate_pdf_markdown_text(text)

            supabase_client.update_summary_job(job_id, {"progress": 30})
            summarized = _summarize_with_chunks_retry(
                client=client,
                text=text,
                job_id=job_id,
            )
            summary_data = summarized["summary_data"]
            final_summary_text = summarized["summary_text"]
            coverage = summarized["coverage"]

        else:
            # DOCX
            _validate_source_text(text)
            summarized = _summarize_with_chunks_retry(
                client=client, text=text, job_id=job_id
            )
            summary_data       = summarized["summary_data"]
            final_summary_text = summarized["summary_text"]
            coverage           = summarized["coverage"]

        # 4. Key points.
        supabase_client.update_summary_job(job_id, {"progress": 88})

        key_points: List[str] = _sanitize_key_points(
            _parse_key_points_from_json(summary_data)
        )

        # Fallback: call KEYPOINTS_SYSTEM_PROMPT if key_points is empty.
        if not key_points and bool(getattr(settings, "SUMMARY_ENABLE_KEYPOINTS_FALLBACK", False)):
            raw_kp = _chat(
                client=client,
                system_prompt=KEYPOINTS_SYSTEM_PROMPT,
                user_prompt=_json.dumps(summary_data, ensure_ascii=False)[:4000],
                max_tokens=int(getattr(settings, "SUMMARY_KEYPOINTS_MAX_TOKENS", "450")),
            )
            try:
                kp_data    = _safe_parse_json(raw_kp)
                key_points = _sanitize_key_points(kp_data.get("key_points", []))
            except Exception:
                # Fallback text parse if JSON parsing fails.
                key_points = _sanitize_key_points(
                    [
                        re.sub(r"^[^\w\[]+\s*", "", ln.strip())
                        for ln in raw_kp.splitlines()
                        if ln.strip()
                    ]
                )

        if not key_points:
            key_points = _sanitize_key_points(_derive_key_points_from_summary_json(summary_data))

        if not key_points:
            key_points = _derive_key_points_from_summary_text(final_summary_text)

        if not key_points and final_summary_text.strip():
            key_points = ["Tom tat da duoc tao, nhung model khong tra ve danh sach y chinh rieng."]

        if not key_points:
            raise RuntimeError("Khong tao duoc tom tat tu noi dung file.")

        final_summary_text = _sanitize_summary_text(final_summary_text)
        if not final_summary_text:
            final_summary_text = _summary_json_to_text(summary_data)
            final_summary_text = _sanitize_summary_text(final_summary_text)
        if not final_summary_text:
            raise RuntimeError("Khong tao duoc summary_text tu noi dung file.")

        summary_data["key_points"] = key_points
        summary_data = _sanitize_summary_json(summary_data)
        key_points = _sanitize_key_points(summary_data.get("key_points", []))

        payload = _build_summary_json_payload(
            job_id=job_id,
            file_name=file_name,
            summary_text=final_summary_text,
            summary_json=summary_data,
            key_points=key_points,
            source_word_count=len(text.split()),
            coverage=coverage,
        )
        _upload_summary_json(
            bucket=bucket,
            user_id=user_id,
            job_id=job_id,
            payload=payload,
        )

        # 5. Save result to DB.
        saved_row, saved_status = supabase_client.update_summary_job(
            job_id,
            {
                "status": "done",
                "progress": 100,
                "summary_text": final_summary_text,
                "summary_json": summary_data,
                "key_points": key_points,
                "source_word_count": len(text.split()),
                "finished_at": now_iso(),
                "error_message": None,
            },
        )
        if saved_status >= 400:
            raise RuntimeError(f"Khong luu duoc tom tat vao Supabase DB: {saved_row}")

    except Exception as exc:
        supabase_client.update_summary_job(
            job_id,
            {
                "status": "failed",
                "progress": 100,
                "finished_at": now_iso(),
                "error_message": str(exc)[:1000] if str(exc) else "Khong ro loi.",
            },
        )
        raise
